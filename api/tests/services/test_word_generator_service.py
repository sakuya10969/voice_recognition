import pytest
from pathlib import Path
import os
from docx import Document
from app.services.word_generator_service import WordGeneratorService

class TestWordGeneratorService:
    @pytest.fixture
    def word_generator(self):
        """WordGeneratorServiceのインスタンスを提供するフィクスチャ"""
        return WordGeneratorService()

    @pytest.fixture
    def sample_texts(self):
        """テスト用のサンプルテキストを提供するフィクスチャ"""
        return {
            'transcribed': 'これはテスト用の文字起こしテキストです。',
            'summarized': 'これはテスト用の要約テキストです。'
        }

    @pytest.mark.asyncio
    async def test_create_word_document_success(self, word_generator, sample_texts):
        """正常系: Word文書生成が成功するケース"""
        # 実行
        file_path = await word_generator.create_word_document(
            sample_texts['transcribed'],
            sample_texts['summarized']
        )

        try:
            # 検証
            assert os.path.exists(file_path)
            doc = Document(file_path)
            
            # 文書構造の検証
            self._verify_document_structure(doc, sample_texts)
        finally:
            # クリーンアップ
            await word_generator.cleanup_word_file(file_path)

    def _verify_document_structure(self, doc: Document, sample_texts: dict):
        """Word文書の構造を検証するヘルパーメソッド"""
        # タイトルの検証
        assert doc.paragraphs[0].text == '文字起こしと要約'
        
        # セクションとコンテンツの検証
        sections = {
            '要約': sample_texts['summarized'],
            '文字起こし': sample_texts['transcribed']
        }
        
        current_section = None
        for para in doc.paragraphs[1:]:  # タイトルをスキップ
            if para.text in sections:
                current_section = para.text
            elif para.text and current_section:
                assert para.text == sections[current_section]

    @pytest.mark.asyncio
    async def test_cleanup_word_file_success(self, word_generator, sample_texts):
        """正常系: ファイルクリーンアップが成功するケース"""
        # 準備
        file_path = await word_generator.create_word_document(
            sample_texts['transcribed'],
            sample_texts['summarized']
        )
        assert os.path.exists(file_path)
        
        # 実行
        await word_generator.cleanup_word_file(file_path)
        
        # 検証
        assert not os.path.exists(file_path)

    @pytest.mark.asyncio
    async def test_create_word_document_error(self, word_generator):
        """異常系: Word文書生成が失敗するケース"""
        with pytest.raises(Exception) as exc_info:
            await word_generator.create_word_document(None, None)
        assert "Word文書の生成に失敗" in str(exc_info.value)

    @pytest.mark.asyncio
    async def test_cleanup_nonexistent_file(self, word_generator):
        """異常系: 存在しないファイルのクリーンアップ"""
        # 実行と検証
        await word_generator.cleanup_word_file("/path/to/nonexistent/file.docx")
        # エラーが発生しないことを確認