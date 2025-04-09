import pytest
import os
from docx import Document
from typing import Dict, Tuple
from pathlib import Path

from app.services.word_generating_service import WordGeneratingService

class TestWordGeneratingService:
    @pytest.fixture
    def word_generating_service(self) -> WordGeneratingService:
        """WordGeneratorServiceのインスタンスを提供するフィクスチャ"""
        return WordGeneratingService()

    @pytest.fixture
    def mock_sample_texts(self) -> Dict[str, str]:
        """テスト用のサンプルテキストを提供するフィクスチャ"""
        return {
            'transcribed_text': 'これはテスト用の文字起こしテキストです。',
            'summarized_text': 'これはテスト用の要約テキストです。'
        }

    async def _create_test_document(
        self,
        word_generating_service: WordGeneratingService,
        mock_sample_texts: Dict[str, str]
    ) -> Tuple[Path, Document]:
        """テスト用のWord文書を作成するヘルパーメソッド"""
        file_path = await word_generating_service.create_word_document(
            mock_sample_texts['transcribed_text'],
            mock_sample_texts['summarized_text']
        )
        doc = Document(file_path)
        return file_path, doc

    async def _verify_document_structure(self, doc: Document, mock_sample_texts: Dict[str, str]) -> None:
        """Word文書の構造を検証するヘルパーメソッド"""
        # タイトルの検証
        assert doc.paragraphs[0].text == '文字起こしと要約'
        
        # セクションとコンテンツの検証
        sections = {
            '要約': mock_sample_texts['summarized_text'],
            '文字起こし': mock_sample_texts['transcribed_text']
        }
        
        current_section = None
        for para in doc.paragraphs[1:]:  # タイトルをスキップ
            if para.text in sections:
                current_section = para.text
            elif para.text and current_section:
                assert para.text == sections[current_section]

    @pytest.mark.asyncio
    async def test_create_word_document_success(
        self, 
        word_generating_service: WordGeneratingService, 
        mock_sample_texts: Dict[str, str],
    ) -> None:
        """正常系: Word文書生成が成功するケース"""
        file_path, doc = await self._create_test_document(word_generating_service, mock_sample_texts)

        try:
            assert os.path.exists(file_path)
            await self._verify_document_structure(doc, mock_sample_texts)
        finally:
            await word_generating_service.cleanup_word_file(str(file_path))

    @pytest.mark.asyncio
    async def test_cleanup_word_file_success(
        self, 
        word_generating_service: WordGeneratingService, 
        mock_sample_texts: Dict[str, str]
    ) -> None:
        """正常系: ファイルクリーンアップが成功するケース"""
        file_path, _ = await self._create_test_document(word_generating_service, mock_sample_texts)
        assert os.path.exists(file_path)
        
        await word_generating_service.cleanup_word_file(str(file_path))
        assert not os.path.exists(file_path)

    @pytest.mark.asyncio
    async def test_create_word_document_error(self, word_generating_service: WordGeneratingService) -> None:
        """異常系: Word文書生成が失敗するケース"""
        with pytest.raises(ValueError, match="文字起こしテキストと要約テキストは必須です"):
            await word_generating_service.create_word_document(None, None)

    @pytest.mark.asyncio
    async def test_cleanup_nonexistent_file(self, word_generating_service: WordGeneratingService) -> None:
        """異常系: 存在しないファイルのクリーンアップ"""
        await word_generating_service.cleanup_word_file("/path/to/nonexistent/file.docx")
        # エラーが発生しないことを確認