import logging
from typing import Optional
from pathlib import Path
import tempfile
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

class WordGeneratorService:
    """Word文書の生成を行うサービス"""
    
    @staticmethod
    async def create_word_document(
        transcribed_text: str,
        summarized_text: str
    ) -> str:
        """文字起こしと要約からWord文書を生成する"""
        try:
            # 一時ディレクトリの作成
            temp_dir = Path(tempfile.gettempdir()) / "voice_recognition"
            temp_dir.mkdir(exist_ok=True)
            
            # 一時ファイルのパスを生成
            temp_file_path = temp_dir / "transcription.docx"
            
            # Word文書の作成
            doc = Document()
            
            # タイトルの追加
            title = doc.add_heading('文字起こしと要約', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 要約セクション
            doc.add_heading('要約', level=1)
            summary_para = doc.add_paragraph(summarized_text)
            summary_para.font.size = Pt(11)
            
            # 文字起こしセクション
            doc.add_heading('文字起こし', level=1)
            trans_para = doc.add_paragraph(transcribed_text)
            trans_para.font.size = Pt(11)
            
            # 文書の保存
            doc.save(str(temp_file_path))
            
            logger.info(f"Word文書を生成: {temp_file_path}")
            return str(temp_file_path)
            
        except Exception as e:
            logger.error(f"Word文書の生成に失敗: {str(e)}")
            raise Exception(f"Word文書の生成に失敗しました: {str(e)}")
    
    @staticmethod
    async def cleanup_word_file(file_path: str) -> None:
        """生成したWord文書を削除する"""
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                logger.info(f"Word文書を削除: {file_path}")
        except Exception as e:
            logger.warning(f"Word文書の削除に失敗: {str(e)}") 