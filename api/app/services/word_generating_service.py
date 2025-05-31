import logging
from pathlib import Path
import tempfile
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

logger = logging.getLogger(__name__)


class WordGeneratingService:
    """Word文書の生成を行うサービス"""

    def __init__(self):
        self.doc = None
        self.temp_dir = None
        self.temp_file_path = None

    async def create_word_document(
        self, transcribed_text: str, summarized_text: str
    ) -> Path:
        """文字起こしと要約からWord文書を生成する"""
        if not transcribed_text or not summarized_text:
            raise ValueError("文字起こしテキストと要約テキストは必須です")
        try:
            self._initialize_document()
            self._add_sections(summarized_text, transcribed_text)
            self._save_document()
            logger.info(f"Word文書を生成: {self.temp_file_path}")
            return self.temp_file_path
        except Exception as e:
            logger.error(f"Word文書の生成に失敗: {str(e)}")
            raise Exception(f"Word文書の生成に失敗しました: {str(e)}")

    def _initialize_document(self) -> None:
        """一時ディレクトリとWord文書の初期化"""
        self.temp_dir = Path(tempfile.mkdtemp())
        now = datetime.now()
        filename = f"{now.year}_{now.month:02d}{now.day:02d}_{now.hour:02d}{now.minute:02d}_議事録.docx"
        self.temp_file_path = self.temp_dir / filename
        self.doc = Document()

    def _add_sections(self, summary_text: str, transcribed_text: str) -> None:
        """タイトル・要約・文字起こしセクションをまとめて追加"""
        self._add_title()
        self._add_section("要約", summary_text)
        self._add_section("文字起こし", transcribed_text)

    def _add_title(self) -> None:
        """タイトルを追加"""
        title = self.doc.add_heading("文字起こしと要約", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_section(self, heading: str, content: str) -> None:
        """見出しと本文を追加"""
        self.doc.add_heading(heading, level=1)
        para = self.doc.add_paragraph(content)
        for run in para.runs:
            run.font.size = Pt(11)

    def _save_document(self) -> None:
        """文書を保存"""
        self.doc.save(str(self.temp_file_path))

    async def cleanup_word_file(self, file_path: str) -> None:
        """生成したWord文書を削除する"""
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                logger.info(f"Word文書を削除: {file_path}")
        except Exception as e:
            logger.warning(f"Word文書の削除に失敗: {str(e)}")
