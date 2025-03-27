import os
from pathlib import Path
from datetime import datetime
import tempfile
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from fastapi import HTTPException

async def create_word(transcribed_text: str, summarized_text: str) -> Path:
    """議事録を作成して一時ファイルパスを返す関数（レイアウト修正版）"""
    try:
        file_name = generate_file_name()
        temp_path = get_temp_path(file_name)
        document = Document()
        add_main_title(document, "議事録")
        add_section(document, "■ 要約結果", summarized_text)
        add_section(document, "■ 文字起こし結果", transcribed_text)
        document.save(temp_path)
        return temp_path
    except Exception as e:
        raise HTTPException(
            status_code=500, detail=f"Failed to create word file: {str(e)}"
        )

def generate_file_name() -> str:
    """ファイル名を生成"""
    return f"議事録_{datetime.now().strftime('%Y-%m%d-%H%M')}.docx"

def get_temp_path(file_name: str) -> Path:
    """一時ファイルのパスを取得"""
    temp_dir = tempfile.gettempdir()
    return Path(temp_dir) / file_name

def add_main_title(document: Document, title_text: str):
    """メインタイトルを追加"""
    title = document.add_heading(title_text, level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def add_section(document: Document, heading: str, text: str, chunk_size: int = 1300):
    """セクションを追加"""
    document.add_heading(heading, level=2)
    chunks = split_text_into_chunks(text, chunk_size)
    for i, chunk in enumerate(chunks):
        paragraph = document.add_paragraph(chunk)
        format_paragraph(paragraph)
        if i < len(chunks) - 1:
            paragraph.add_run("\n")
            document.add_page_break()

def split_text_into_chunks(text: str, chunk_size: int) -> list:
    """テキストをチャンクに分割"""
    return [text[i : i + chunk_size] for i in range(0, len(text), chunk_size)]

def format_paragraph(paragraph):
    """段落のフォーマットを設定"""
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    paragraph.runs[0].font.size = Pt(11)

async def cleanup_word(file_path: str):
    """一時ファイルを削除"""
    try:
        if os.path.exists(file_path):
            os.remove(file_path)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to cleanup file: {str(e)}")
